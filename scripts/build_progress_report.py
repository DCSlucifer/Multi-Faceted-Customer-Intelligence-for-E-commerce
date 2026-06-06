"""Generate the Final Project *Progress* Report as a Word document.

Deliberately written in a mid-project, planning tone (problem + dataset +
*intended* methodology + progress so far), per the assignment brief. It does NOT
claim the project is finished.

    python scripts/build_progress_report.py   ->  reports/progress_report.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "progress_report.docx"

PRIMARY = RGBColor(0x26, 0x46, 0x53)


def main() -> int:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def heading(text: str) -> None:
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = PRIMARY
            run.font.size = Pt(14)

    def para(text: str, italic: bool = False) -> None:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        if italic:
            for r in p.runs:
                r.font.italic = True

    def bullet(text: str) -> None:
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(2)

    # ---------------- Title block ---------------- #
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Final Project — Progress Report")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Multi-Faceted Customer Intelligence for E-commerce\n"
                     "An End-to-End Data Mining Pipeline on Online Retail II")
    rs.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        "Data Mining 2026 — Team 4, University of Information Technology, VNU-HCM\n"
        "Nguyen Minh Cuong (22520177) · Vo Thanh Danh (22520201) · "
        "Nguyen Vinh Dat (22520228) · Nguyen Huu Dinh (22520251)")
    rm.font.size = Pt(10)
    rm.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # ---------------- 1. Problem definition ---------------- #
    heading("1. Problem Definition")
    para(
        "Online retailers routinely ask three separate questions about their "
        "customer base: who their customers are (segmentation), which customers "
        "are about to stop buying (churn), and which products tend to be bought "
        "together (market-basket patterns). In practice these questions are usually "
        "tackled in isolation, which misses the connections between them — for "
        "example, a retention campaign is far more effective when the offer matches "
        "the purchasing pattern of the customer's segment.")
    para(
        "For our final project we plan to build a single, coherent data-mining "
        "pipeline that addresses all three questions on the same dataset and then "
        "combines their outputs into a cross-task analysis (e.g., churn rate per "
        "segment and the basket patterns that characterise each segment). The "
        "practical motivation is well established: retaining an existing customer is "
        "substantially cheaper than acquiring a new one, so being able to identify "
        "which customers will churn — and why, in the context of their segment and "
        "baskets — supports concrete, targeted retention decisions.")
    para(
        "Concretely, the project aims to (i) segment customers from their purchase "
        "behaviour, (ii) predict 90-day churn, (iii) mine product association rules, "
        "and (iv) synthesise these results into segment-aware business "
        "recommendations.")

    # ---------------- 2. Dataset ---------------- #
    heading("2. Dataset")
    para(
        "We use the Online Retail II dataset from the UCI Machine Learning "
        "Repository (dataset ID 502, "
        "https://archive.ics.uci.edu/dataset/502/online+retail+ii). It contains "
        "real transaction records from a UK-based online retailer between December "
        "2009 and December 2011 — about 1.07 million invoice-line records, roughly "
        "5,900 unique customers and 5,300 stock items, spread across 41 countries "
        "(predominantly the United Kingdom).")
    para("Each record has the following main fields:")
    bullet("Invoice — transaction/invoice number (a leading 'C' marks a cancellation).")
    bullet("StockCode and Description — product identifier and text label.")
    bullet("Quantity and Price — units purchased and unit price.")
    bullet("InvoiceDate — timestamp of the transaction.")
    bullet("Customer ID — anonymised customer identifier (missing on ~25% of rows).")
    bullet("Country — customer's country.")
    para(
        "From these raw fields we plan to engineer per-customer features: the "
        "classic RFM variables (Recency, Frequency, Monetary) together with simple "
        "behavioural features such as average basket value, average basket size, "
        "number of unique products purchased, and the customer's dominant country. "
        "These engineered features will be the common input shared across the "
        "segmentation and churn tasks.")

    # ---------------- 3. Intended methodology ---------------- #
    heading("3. Intended Methodology")
    para(
        "Our planned pipeline covers preprocessing and four mining tasks; the main "
        "steps are outlined briefly below.")

    para("Preprocessing and feature engineering.", italic=True)
    para(
        "We will clean the data by removing rows with missing Customer ID, dropping "
        "cancellation invoices and non-positive Quantity/Price values, and "
        "winsorising extreme Quantity/Price values. We then derive the RFM and "
        "behavioural features above, and construct a churn label using a temporal "
        "cutoff date with a fixed 90-day post-cutoff observation window (a customer "
        "active before the cutoff but absent during the window is labelled churned).")

    para("Customer segmentation (clustering).", italic=True)
    para(
        "We intend to compare three clustering families — K-means (partitional), "
        "DBSCAN (density-based), and agglomerative/Ward hierarchical clustering — on "
        "log-transformed, standardised RFM features. Cluster quality will be assessed "
        "with internal validity indices (Silhouette, Davies-Bouldin, "
        "Calinski-Harabasz), and the segments will be visualised in 2-D using PCA "
        "and t-SNE. We also plan to check that the chosen segmentation is stable "
        "when K-means is re-run across several random seeds.")

    para("Churn classification.", italic=True)
    para(
        "We plan to train several classical classifiers (e.g., Logistic Regression, "
        "Decision Tree, Random Forest, k-NN, SVM, and Naive Bayes) and a small "
        "neural network (MLP), using stratified cross-validation plus a held-out "
        "test set, and addressing class imbalance with SMOTE applied to the training "
        "data only. Models will be compared on AUC, F1, precision and recall. To put "
        "those scores in context we also intend to evaluate simple baselines (a "
        "majority-class predictor and a recency-only model), run a feature-group "
        "ablation, and check the calibration of the predicted probabilities. A key "
        "design concern we are paying attention to is avoiding target leakage when "
        "building the churn features (computing them only from data before the cutoff).")

    para("Association rule mining.", italic=True)
    para(
        "We will mine frequent itemsets and association rules over invoice-level "
        "baskets — restricted to frequently-purchased items to keep mining "
        "tractable — using both the Apriori and FP-Growth algorithms, ranking rules "
        "by support, confidence and lift, and comparing the two algorithms' "
        "runtimes. We will also mine rules within individual segments for the "
        "cross-task analysis.")

    para("Deep-learning extension and cross-task synthesis.", italic=True)
    para(
        "As an extension we plan to add an autoencoder for unsupervised detection of "
        "anomalous customers (via reconstruction error) and to use SHAP to interpret "
        "the churn model's drivers. Finally, the cross-task step will join the "
        "segmentation, churn, and association-rule results to produce per-segment "
        "churn rates and per-segment basket patterns — the analysis we expect to "
        "yield the most useful business insight.")

    para(
        "Throughout, we aim to keep the work reproducible (fixed random seed, a "
        "documented run procedure, and shared utility code) so that all reported "
        "numbers can be regenerated.")

    # ---------------- 4. Progress + plan ---------------- #
    heading("4. Progress to Date and Remaining Work")
    para("Completed so far:")
    bullet("Project setup: repository structure, environment, and shared utilities.")
    bullet("Dataset acquired and loaded; initial cleaning and exploratory data analysis carried out (the cleaning step reduces the ~1.07M raw rows to roughly 0.8M usable transactions).")
    bullet("RFM and behavioural feature engineering implemented; first clustering experiments run.")
    para("In progress:")
    bullet("Churn classification (classical models and the MLP) and association-rule mining.")
    bullet("Refining the churn feature construction to ensure it is leakage-free.")
    para("Remaining work before the final submission:")
    bullet("Deep-learning extension (autoencoder + SHAP) and the cross-task synthesis.")
    bullet("Consolidating evaluation tables and figures, then writing the final IEEE-format report and slide deck.")

    doc.save(OUT)
    print(f"Wrote {OUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
